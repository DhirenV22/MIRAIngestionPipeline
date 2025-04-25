import fitz  # PyMuPDF
from docx import Document


def extract_pdf_text(file_path):
    """
    Extracts text from a PDF file using PyMuPDF.

    :param file_path: Path to the PDF file.
    :return: Extracted text as a string.
    """
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        print(f"Extracted text from PDF: {file_path}")
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text


def extract_docx_text(file_path):
    """
    Extracts text from a DOCX file, including paragraphs and tables.

    :param file_path: Path to the DOCX file.
    :return: Extracted text as a string.
    """
    text = ""
    try:
        doc = Document(file_path)
        # Extract text from paragraphs
        text += "\n".join([para.text for para in doc.paragraphs])

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text += "\n" + "\t".join(row_text)
        print(f"Extracted text from DOCX: {file_path}")
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text
